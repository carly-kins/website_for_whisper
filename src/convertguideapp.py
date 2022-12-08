import whisper
import cv2
import os
import glob
from docx import Document
from docx.shared import Inches

def clean_path(dir):
    for f in os.listdir(dir):
        os.remove(os.path.join(dir, f))

def convert_guide(model_type, screenshots, dir, video, doc):
    model = whisper.load_model(model_type)
    result = model.transcribe(audio=video, fp16=False)
    segments = result['segments']
    document = Document()
    document.add_heading('Whisper Generated Guide', 0)

    seconds = []
    for item in range(len(segments)):
        seconds.append(round(segments[item]['start']))

    averages = []
    total = 0
    count = 0

    for sec in seconds:
        total += sec
        count += 1
        if count == screenshots:
            average = round(total / screenshots)
            averages.append(average)

            total = 0
            count = 0

    if count > 0:
        average = total / count
        averages.append(average) 

    lines = []
    for item in range(len(segments)):
        lines.append(segments[item]['text'])
    
    combine_lines = []
    temp = ''
    count = 0
    
    for line in lines:
        temp += line
        count += 1
        if count == screenshots:
            combine_lines.append(temp)

            temp = ''
            count = 0

    if count > 0:
        combine_lines.append(temp)

    vidcap = cv2.VideoCapture(video)
    success,image = vidcap.read()
    success = True
    while success:
        for avg in averages: 
            vidcap.set(cv2.CAP_PROP_POS_MSEC,(avg*1000))
            success,image = vidcap.read()
            cv2.imwrite(dir + str(avg) + ".jpg", image)  
        break
    else:
        print("Error Saving Images")

    for second, line in zip(averages, combine_lines):
        document.add_picture(dir + str(second) + ".jpg", width=Inches(3))
        document.add_paragraph(line)

    document.save(dir + doc)
    imageFiles = glob.glob(os.path.join(dir, "*.jpg"))
    for image in imageFiles:
        os.remove(image)