import io
import asyncio

from fastapi import FastAPI, Request, File, Query
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates

import whisper
import cv2
import os
import glob
from docx import Document
from docx.shared import Inches

app = FastAPI()

app.mount('/static', StaticFiles(directory='static'), name='static')
template = Jinja2Templates(directory='templates')

@app.get('/', response_class=HTMLResponse)
def index(request: Request):
    return template.TemplateResponse('index.html',  {"request": request, "disabled": True, "result": False})

@app.post('/')
def convert_guide(request: Request, file: bytes = File()):
    with open('audio.mp4', 'wb') as f:
        f.write(file)
    
    path = "static"
    imageFiles = glob.glob(os.path.join(path, "*.jpg"))
    docFiles = glob.glob(os.path.join(path, "*.docx"))
    for image in imageFiles:
        os.remove(image)
    for doc in docFiles:
        os.remove(doc)

    data = request.form()
    model_type = asyncio.run(data)['model_type']
    model = whisper.load_model(model_type)
    result = model.transcribe(audio="audio.mp4", fp16=False)
    segments = result['segments']

    seconds = []
    for item in range(len(segments)):
        seconds.append(round(segments[item]['start']))
    lines = []
    for item in range(len(segments)):
        lines.append(segments[item]['text'])

    averages = []
    total = 0
    count = 0

    for sec in seconds:
        total += sec
        count += 1
        if count == 3:
            average = round(total / 3)
            averages.append(average)

            total = 0
            count = 0

    if count > 0:
        average = total / count
        averages.append(average)  

    combine_lines = []
    temp = ''
    count = 0
    
    for line in lines:
        temp += line
        count += 1
        if count == 3:
            combine_lines.append(temp)

            temp = ''
            count = 0

    if count > 0:
        combine_lines.append(temp) 

    vidcap = cv2.VideoCapture("audio.mp4")
    success,image = vidcap.read()
    success = True
    while success:
        for avg in averages: 
            vidcap.set(cv2.CAP_PROP_POS_MSEC,(avg*1000))
            success,image = vidcap.read()
            cv2.imwrite("static/" + str(avg) + ".jpg", image)  
        break
    else:
        print("Error Saving Images")

    document = Document()

    document.add_heading('Whisper Generated Guide', 0)

    for second, line in zip(averages, combine_lines):
        document.add_picture("static/" + str(second) + ".jpg", width=Inches(2))
        document.add_paragraph(line)

    document.save('static/output.docx')

    return template.TemplateResponse('index.html',  {"request": request, "disabled": False, "result": True})